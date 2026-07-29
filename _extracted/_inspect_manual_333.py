import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"

docx = Path(r"C:\Users\user\Desktop\333\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_ispravleno.docx")
folder = Path(r"C:\Users\user\Desktop\333")
out = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\manual_333_status.txt")
marker = "\u041a\u043e\u043d\u0435\u0446 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430"

files = []
for p in sorted(folder.iterdir()):
    files.append((p.name, p.stat().st_size, "dir" if p.is_dir() else "file"))

with zipfile.ZipFile(docx) as z:
    root = ET.fromstring(z.read("word/document.xml"))

has_pgBorders = "yes" if root.findall(".//" + W + "pgBorders") else "no"
inlines = root.findall(".//" + WP + "inline")
num_inline = len(inlines)
cxs = [ext.get("cx") for ext in root.findall(".//" + WP + "extent") if ext.get("cx") is not None]
unique_cxs = sorted(set(cxs), key=lambda x: int(x))

all_sect = root.findall(".//" + W + "sectPr")
sect = all_sect[-1] if all_sect else None
pgSz = sect.find(W + "pgSz") if sect is not None else None
pgMar = sect.find(W + "pgMar") if sect is not None else None

def local_attrs(el):
    if el is None:
        return {}
    return {k.split("}")[-1]: v for k, v in el.attrib.items()}

page_size = local_attrs(pgSz)
margins = local_attrs(pgMar)

paras = root.findall(".//" + W + "p")
para_count = len(paras)

def para_text(p):
    return "".join((t.text or "") for t in p.findall(".//" + W + "t")).strip()

nonempty = [t for t in (para_text(p) for p in paras) if t]
last5 = nonempty[-5:]
has_konec = any(marker in t for t in nonempty)

docx_nonempty = []
docx_block = []
try:
    from docx import Document
    d = Document(str(docx))
    docx_nonempty = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    sec = d.sections[-1]
    konec_docx = "yes" if any(marker in t for t in docx_nonempty) else "no"
    docx_block.append("")
    docx_block.append("python-docx:")
    docx_block.append("  paragraphs: %d" % len(d.paragraphs))
    docx_block.append("  nonempty paragraphs: %d" % len(docx_nonempty))
    docx_block.append("  inline_shapes: %d" % len(d.inline_shapes))
    docx_block.append("  page width/height (inches): %.3f x %.3f" % (sec.page_width.inches, sec.page_height.inches))
    docx_block.append("  margins inches L/R/T/B: %.3f/%.3f/%.3f/%.3f" % (
        sec.left_margin.inches, sec.right_margin.inches, sec.top_margin.inches, sec.bottom_margin.inches))
    docx_block.append("  last 5 nonempty:")
    for t in docx_nonempty[-5:]:
        docx_block.append("    - %s" % t[:200])
    docx_block.append("  Konec dokumenta: %s" % konec_docx)
    docx_block.append("")
except Exception as e:
    docx_block = ["", "python-docx: ERROR %s" % e, ""]

lines = [
    "manual_333_status",
    "=" * 60,
    "",
    r"1. Folder: C:\Users\user\Desktop\333\\",
]
for name, size, kind in files:
    lines.append("   %s\t%d bytes\t(%s)" % (name, size, kind))

cx_in = ["%s EMU (~%.3f in)" % (c, int(c) / 914400) for c in unique_cxs]
lines += [
    "",
    "2. DOCX checks (zipfile / python-docx)",
    "   File: %s" % docx.name,
    "   Size: %d bytes" % docx.stat().st_size,
    "   Has pgBorders: %s" % has_pgBorders,
    "   Inline images (wp:inline): %d" % num_inline,
    "   wp:extent count: %d" % len(cxs),
    "   Unique wp:extent cx values: %s" % unique_cxs,
    "   Unique cx (with inches): %s" % cx_in,
    "   Page size (w:pgSz): %s" % page_size,
]
if page_size.get("w") and page_size.get("h"):
    lines.append("   Page size inches: %.3f x %.3f (A4 portrait)" % (int(page_size["w"]) / 1440, int(page_size["h"]) / 1440))
lines.append("   Margins (w:pgMar): %s" % margins)
for k in ("left", "right", "top", "bottom", "header", "footer", "gutter"):
    if k in margins:
        lines.append("     %s: %s twips (~%.3f in)" % (k, margins[k], int(margins[k]) / 1440))
lines += [
    "   Approx paragraph count (w:p in document.xml): %d" % para_count,
    "   Non-empty paragraphs: %d" % len(nonempty),
    "   '%s' present: %s" % (marker, "yes" if has_konec else "no"),
    "   Last 5 non-empty paragraphs (document.xml):",
]
for t in last5:
    lines.append("     - %s" % t[:300])
lines.extend(docx_block)

text = "\n".join(lines)
out.parent.mkdir(parents=True, exist_ok=True)
out.write_text(text, encoding="utf-8")
print("Wrote:", out)
print("pgBorders:", has_pgBorders)
print("inline:", num_inline)
print("unique_cx:", unique_cxs)
print("has_konec:", has_konec)
for t in last5:
    print("LAST:", t)
for t in docx_nonempty[-5:]:
    print("DOCX:", t)