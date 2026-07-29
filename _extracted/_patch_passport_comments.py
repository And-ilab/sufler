# -*- coding: utf-8 -*-
"""Apply Word comments to GW-DS09 passport; save clean file to Desktop/333."""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

SRC = Path(r"c:\Users\user\Downloads\Паспорт_станка_GW_DS09_ГОСТ_2_601_2019_1.docx")
OUT = Path(r"C:\Users\user\Desktop\333\Pasport_stanka_GW_DS09_GOST_2_601_2019.docx")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

SUPPLIER_LINES = [
    "ПОСТАВЩИК:",
    "ООО «Финсельват»",
    "Адрес: Минская обл., Минский р-н, Новодворский сельсовет,",
    "д. Большое Стиклево, д. 40, к. 2, оф. 52",
    "УНП 692204462",
    "р/с BY21ALFA30122C18740010270000 в ЗАО «Альфа-банк»,",
    "г. Минск, ул. Сурганова, 43–47, код ALFABY2X",
    "e-mail: finselvat.info@yandex.ru; www.цифровая.бел",
    "",
    "Официальный представитель и сервисный центр на территории РБ:",
    "тел. +375 29 667 88 73",
    "HelpDesk@digitranslab.com",
    "",
    "Приложение к договору поставки № 3004;",
    "согласованная КД к договору № 2802/1 от 28.02.2024",
]

# media filename -> caption (only real figures from comments + packaging photo)
CAPTIONS = {
    "image2.png": "Рис. 1. Общий вид двухстанционной намоточной машины GW-DS09",
    "image4.png": "Рис. 2. Габаритные размеры станка GW-DS09",
    "image1.png": "Рис. 3. Компоновка рабочих станций (3D-модель)",
    "image3.png": "Рис. 4. Упаковка изделия и заводская табличка",
}

# Max display width for figures (cm) — keep inside A4 frame
MAX_FIG_WIDTH_CM = {
    "image2.png": 9.0,   # was overflowing page height
    "image4.png": 14.0,
    "image1.png": 8.0,
    "image3.png": 12.0,
}


def set_run_font(run, size=11, bold=False, name="Times New Roman"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text_multiline(cell, lines, *, first_bold=True, size=9):
    """Replace cell content with multiple paragraphs."""
    # clear existing paragraphs except keep one
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    for i, line in enumerate(lines):
        p = OxmlElement("w:p")
        tc.append(p)
        from docx.text.paragraph import Paragraph

        para = Paragraph(p, cell)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0 if line else 2)
        para.paragraph_format.line_spacing = 1.0
        if line:
            run = para.add_run(line)
            set_run_font(run, size=size, bold=(first_bold and i == 0))


def ensure_complete_borders(table, color="000000", sz="8"):
    """Comment 2–6: finish table frame — all four edges on every cell + tblBorders."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # remove old tblBorders
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)

    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag == qn("w:tcBorders"):
                    tcPr.remove(child)
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), sz)
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)


def rels_target_map(docx_path: Path) -> dict[str, str]:
    """rId -> target filename (e.g. media/image2.png)."""
    with zipfile.ZipFile(docx_path) as z:
        root = ET.fromstring(z.read("word/_rels/document.xml.rels"))
    out = {}
    for rel in root:
        rid = rel.get("Id")
        target = rel.get("Target", "")
        if rid:
            out[rid] = target
    return out


def find_drawing_paragraphs(doc: Document, rid_map: dict[str, str]):
    """Yield (paragraph, media_basename) for each blip image."""
    results = []
    for p in doc.paragraphs:
        for blip in p._element.iter(A + "blip"):
            rid = blip.get(qn("r:embed")) or blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not rid:
                continue
            target = rid_map.get(rid, "")
            name = Path(target).name
            if name:
                results.append((p, name))
    return results


def resize_image_in_paragraph(p, max_width_cm: float):
    target_cx = int(Cm(max_width_cm))
    for extent in p._element.iter(WP + "extent"):
        cx = int(extent.get("cx", "0") or 0)
        cy = int(extent.get("cy", "0") or 0)
        if cx <= 0:
            continue
        # only shrink, don't enlarge small images beyond limit
        if cx <= target_cx:
            # still cap very tall images: if height > 18cm, scale by height
            max_cy = int(Cm(16.0))
            if cy > max_cy:
                scale = max_cy / cy
                extent.set("cx", str(int(cx * scale)))
                extent.set("cy", str(max_cy))
            continue
        new_cy = int(cy * (target_cx / cx)) if cy else cy
        # also cap height
        max_cy = int(Cm(16.0))
        if new_cy > max_cy:
            scale = max_cy / new_cy
            target_cx = int(target_cx * scale)
            new_cy = max_cy
        extent.set("cx", str(target_cx))
        extent.set("cy", str(new_cy))
    for ext in p._element.iter(A + "ext"):
        cx = ext.get("cx")
        cy = ext.get("cy")
        if not cx or not cy:
            continue
        try:
            icx, icy = int(cx), int(cy)
        except ValueError:
            continue
        if icx < 50000:
            continue
        # match nearest wp:extent
        for extent in p._element.iter(WP + "extent"):
            ext.set("cx", extent.get("cx"))
            ext.set("cy", extent.get("cy"))
            break


def insert_caption_after(paragraph, text: str):
    from docx.text.paragraph import Paragraph

    # skip if next sibling already has this caption
    nxt = paragraph._element.getnext()
    if nxt is not None and nxt.tag == qn("w:p"):
        existing = "".join(t.text or "" for t in nxt.iter(W + "t"))
        if existing.strip().startswith("Рис."):
            return Paragraph(nxt, paragraph._parent)

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=False)
    return p


def setup_page_and_footer(doc: Document):
    """A4 margins only; no page frame (pgBorders) and no footer caption/PAGE field."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    # remove page frame if present; do not add pgBorders
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgBorders"):
            sectPr.remove(child)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    for child in list(footer._element):
        if child.tag == qn("w:tbl"):
            footer._element.remove(child)
    # leave empty footer (no caption / PAGE field)
    footer.add_paragraph()


def ensure_section1_page_break(doc: Document):
    """Start section 1 heading on a new page."""
    heading = "1. Основные сведения об изделии"
    for p in doc.paragraphs:
        if heading in (p.text or ""):
            p.paragraph_format.page_break_before = True
            return True
    return False


def strip_comments(path: Path):
    tmp = path.with_suffix(".tmp.docx")
    skip = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsExtensible.xml",
        "word/commentsIds.xml",
        "word/people.xml",
    }
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            if name in skip:
                continue
            data = zin.read(name)
            if name == "word/document.xml":
                root = ET.fromstring(data)
                for el in list(root.iter()):
                    for child in list(el):
                        if child.tag in (
                            W + "commentRangeStart",
                            W + "commentRangeEnd",
                            W + "commentReference",
                        ):
                            el.remove(child)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                for rel in list(root):
                    target = (rel.get("Target") or "").lower()
                    if "comment" in target or target.endswith("people.xml"):
                        root.remove(rel)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "[Content_Types].xml":
                root = ET.fromstring(data)
                for el in list(root):
                    pn = (el.get("PartName") or "").lower()
                    if "comment" in pn or pn.endswith("/people.xml"):
                        root.remove(el)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(name, data)
    tmp.replace(path)


def main():
    if not SRC.exists():
        raise SystemExit(f"Missing source: {SRC}")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, OUT)

    # Need rid map from the copy before heavy edits
    rid_map = rels_target_map(OUT)
    doc = Document(str(OUT))

    # --- Comment 0: supplier block ---
    t0 = doc.tables[0]
    set_cell_text_multiline(t0.cell(0, 1), SUPPLIER_LINES, first_bold=True, size=9)
    print("Comment 0: supplier block updated")

    # --- Comments 2–6 (+0/9): complete table frames ---
    for i, table in enumerate(doc.tables):
        ensure_complete_borders(table, color="000000", sz="8")
    print(f"Comments 2–6: borders completed on {len(doc.tables)} tables")

    # --- Page setup: A4 margins, no frame/footer caption; section 1 on new page ---
    setup_page_and_footer(doc)
    if ensure_section1_page_break(doc):
        print("Section 1: page_break_before set")
    else:
        print("WARN: section 1 heading not found for page break")
    print("Comment 1: A4 margins only (no pgBorders / footer caption)")

    # --- Captions + resize figures (comments 1, 7, 8 + image3) ---
    drawings = find_drawing_paragraphs(doc, rid_map)
    captioned = set()
    for p, name in drawings:
        if name not in CAPTIONS:
            continue
        if name in MAX_FIG_WIDTH_CM:
            resize_image_in_paragraph(p, MAX_FIG_WIDTH_CM[name])
        # special: image2 was too tall — force height cap again after width
        if name == "image2.png":
            resize_image_in_paragraph(p, 8.5)
        if name not in captioned:
            insert_caption_after(p, CAPTIONS[name])
            captioned.add(name)
            print(f"Captioned {name}: {CAPTIONS[name]}")

    doc.save(str(OUT))
    strip_comments(OUT)

    # verify
    with zipfile.ZipFile(OUT) as z:
        names = z.namelist()
        has_comments = "word/comments.xml" in names
        docxml = z.read("word/document.xml")
        has_footer = any(n.startswith("word/footer") for n in names)
        has_borders = docxml.count(b"pgBorders")  # expect 0
        has_finsel = "Финсельват".encode("utf-8") in docxml
        has_helpdesk = b"HelpDesk@digitranslab.com" in docxml
        has_buyer = "ПОКУПАТЕЛЬ".encode("utf-8") in docxml
        fig_count = docxml.count("Рис.".encode("utf-8"))
    print(
        f"Saved {OUT} ({OUT.stat().st_size} bytes)\n"
        f"  comments.xml={has_comments} footer={has_footer} pgBorders_count={has_borders}\n"
        f"  Finselvat={has_finsel} HelpDesk={has_helpdesk} ПОКУПАТЕЛЬ_left={has_buyer}\n"
        f"  'Рис.' occurrences={fig_count}"
    )


if __name__ == "__main__":
    main()
