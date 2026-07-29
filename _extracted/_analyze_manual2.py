# -*- coding: utf-8 -*-
import re, zipfile, sys
from pathlib import Path
from xml.etree import ElementTree as ET
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

DOCX = Path(r"c:\Users\user\Downloads\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina1.docx")
OUT = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\manual_structure.txt")
SUM = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\_manual_findings_summary.txt")
REL = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\_related_files.txt")
SLIDES_SUM = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\_slides_captions_summary.txt")

doc = Document(str(DOCX))
paras = [(i, p.style.name if p.style else "", p.text) for i, p in enumerate(doc.paragraphs)]

starts = [i for i, s, t in paras if re.match(r"^1\.1\.4\b", t.strip())]
ends_candidates = {}
for s in starts:
    end = None
    for i, st, t in paras:
        if i <= s:
            continue
        if re.match(r"^1\.1\.5\b", t.strip()) or re.match(r"^1\.2\b", t.strip()):
            end = i
            break
    ends_candidates[s] = end if end else min(len(paras), s + 20)

extra = []
extra.append("\n" + "=" * 80)
extra.append("## 7b. SECTION 1.1.4 FULL BODY TEXT (corrected)")
for s, e in ends_candidates.items():
    extra.append(f"--- start_para={s} end_para={e} ---")
    for i, st, t in paras[max(0, s - 2) : e]:
        extra.append(f"[{i}] {st}: {t}")

extra.append("\n## 6b. LANGUAGE HITS (refined)")
for i, st, t in paras:
    if re.search(r"китайск|английск|русск|Chinese|English|Russian", t, re.I):
        extra.append(f"[{i}]: {t.strip()}")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if re.search(r"китайск|английск|русск|Chinese|English|Russian", cell.text, re.I):
                extra.append(f"table[{ti}]R{ri}C{ci}: {cell.text.strip()}")

extra.append("\n## 8b. ООО / COMPANY NAME HITS (refined)")
pat = re.compile(
    r"ООО\b|АО\b|ЗАО\b|Finselvat|Финсельват|изготовитель|производитель|поставщик|manufacturer|supplier",
    re.I,
)
hits = []
for i, st, t in paras:
    if pat.search(t):
        hits.append(f"[{i}]: {t.strip()}")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        ct = " | ".join(c.text.strip() for c in row.cells)
        if pat.search(ct):
            hits.append(f"table[{ti}]R{ri}: {ct}")
for si, sec in enumerate(doc.sections):
    for kind, part in [("header", sec.header), ("footer", sec.footer)]:
        for p in part.paragraphs:
            if pat.search(p.text):
                hits.append(f"sec{si}.{kind}: {p.text.strip()}")
extra.extend(hits if hits else ["(no ООО/Finselvat/изготовитель company strings found)"])

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}
extra.append("\n## 2b. DRAWING NODES DETAIL")
with zipfile.ZipFile(DOCX) as z:
    root = ET.fromstring(z.read("word/document.xml"))
    body = root.find("w:body", NS)
    p_idx = -1
    draw_count = 0
    for child in list(body):
        if child.tag == "{%s}p" % NS["w"]:
            p_idx += 1
            drawings = child.findall(".//w:drawing", NS)
            picts = child.findall(".//w:pict", NS)
            if drawings or picts:
                draw_count += len(drawings) + len(picts)
                kinds = []
                for d in drawings:
                    for dp in d.findall(".//wp:docPr", NS):
                        kinds.append(f"name={dp.get('name')} descr={dp.get('descr')}")
                    if d.find(".//a:blip", NS) is not None:
                        kinds.append("blip-image")
                    else:
                        kinds.append("non-blip-drawing")
                before = paras[p_idx - 1][2].strip()[:120] if p_idx > 0 else ""
                after = paras[p_idx + 1][2].strip()[:120] if p_idx + 1 < len(paras) else ""
                own = paras[p_idx][2].strip()[:120]
                extra.append(f"para[{p_idx}] drawings={len(drawings)} picts={len(picts)} kinds={kinds}")
                extra.append(f"  before: {before!r}")
                extra.append(f"  own: {own!r}")
                extra.append(f"  after: {after!r}")
        elif child.tag == "{%s}tbl" % NS["w"]:
            drawings = child.findall(".//w:drawing", NS)
            if drawings:
                extra.append(f"table after para_idx={p_idx}: drawings={len(drawings)}")
    extra.append(f"Total drawing/pict nodes located: {draw_count}")
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    extra.append(f"word/media count: {len(media)}")

extra.append("\n## 3b. PAGE SETUP QUICK")
for i, sec in enumerate(doc.sections):
    extra.append(
        f"section{i}: {sec.page_width.cm:.2f}x{sec.page_height.cm:.2f} cm "
        f"margins LRTB={sec.left_margin.cm:.2f}/{sec.right_margin.cm:.2f}/{sec.top_margin.cm:.2f}/{sec.bottom_margin.cm:.2f}"
    )

extra.append("\n## 5b. TITLE PAGE PARAS 0-15")
for i, st, t in paras[:16]:
    if t.strip():
        extra.append(f"[{i}]: {t.strip()}")

prev = OUT.read_text(encoding="utf-8")
# avoid double-append
if "## 7b. SECTION 1.1.4" not in prev:
    OUT.write_text(prev + "\n".join(extra) + "\n", encoding="utf-8")
else:
    OUT.write_text(prev.split("## 7b. SECTION 1.1.4")[0].rstrip() + "\n" + "\n".join(extra) + "\n", encoding="utf-8")
print("structure size", OUT.stat().st_size)

# Summary
sl = []
sl.append(f"STRUCTURE FILE WRITTEN: {OUT}")
sl.append(f"paragraphs={len(paras)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
sl.append(f"sections={len(doc.sections)}")
for i, sec in enumerate(doc.sections):
    sl.append(
        f"  sec{i}: {sec.page_width.cm:.2f}x{sec.page_height.cm:.2f} cm "
        f"L={sec.left_margin.cm:.2f} R={sec.right_margin.cm:.2f} T={sec.top_margin.cm:.2f} B={sec.bottom_margin.cm:.2f}"
    )
    ht = " | ".join(p.text.strip() for p in sec.header.paragraphs if p.text.strip()) or "(empty)"
    ft = " | ".join(p.text.strip() for p in sec.footer.paragraphs if p.text.strip()) or "(empty)"
    sl.append(f"  header: {ht}")
    sl.append(f"  footer: {ft}")
sl.append("TITLE:")
for i, st, t in paras[:12]:
    if t.strip():
        sl.append("  " + t.strip())
sl.append("SECTION 1.1.4:")
for s, e in ends_candidates.items():
    role = "TOC" if s < 40 else "BODY"
    sl.append(f"  ({role} start={s} end={e})")
    if role == "BODY":
        for i, st, t in paras[s:e]:
            sl.append(f"  [{i}] {t}")
sl.append("LANG:")
for i, st, t in paras:
    if re.search(r"китайск|английск|русск", t, re.I):
        sl.append(f"  [{i}] {t.strip()}")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if re.search(r"китайск|английск|русск", cell.text, re.I):
                sl.append(f"  table[{ti}] {cell.text.strip()}")
sl.append("COMPANY:")
if hits:
    sl.extend("  " + h for h in hits)
else:
    sl.append("  (none: no ООО / Finselvat / изготовитель / поставщик)")
sl.append("HEADINGS body:")
for i, st, t in paras:
    if re.match(r"^\d+(\.\d+)*\s+\S", t.strip()) and i >= 54:
        sl.append(f"  [{i}] {t.strip()}")
SUM.write_text("\n".join(sl), encoding="utf-8")
print("summary ok")

# Related files search
import os
needles = re.compile(
    r"finselvat|финсельват|namotochn|намоточ|double.?station|winding|паспорт|gw-ds09|parameter settings|unpacking",
    re.I,
)
paths = []
roots = [Path.home() / "Desktop", Path.home() / "Downloads"]
skip_parts = {".venv", "venv", "node_modules", "__pycache__", "site-packages"}

def ok_path(p: Path):
    return not any(part in skip_parts for part in p.parts)

for root in roots:
    if not root.exists():
        continue
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        if not ok_path(p):
            continue
        if needles.search(p.name):
            paths.append(str(p))
            continue
        # passport-like docx/pdf/pptx with gw
        if p.suffix.lower() in {".docx", ".pdf", ".pptx", ".md"} and re.search(
            r"паспорт|GW-DS|намот|winding|double|финс|finsel|parameter|unpack|namot", p.name, re.I
        ):
            paths.append(str(p))

# also _extracted
ex = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted")
for p in ex.iterdir():
    if p.is_file() and needles.search(p.name):
        paths.append(str(p))

paths = sorted(set(paths))
REL.write_text("\n".join(paths) + "\n", encoding="utf-8")
print("related", len(paths))
for p in paths:
    print(p)

# Slides summary
parts = []
slide_files = [
    Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\Double_station_winding_machine_unpacking_steps_+_threading_operation_slides.txt"),
    Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\Parameter settings (2)_slides.txt"),
]
for f in slide_files:
    parts.append("=" * 70)
    parts.append(f.name)
    if f.exists():
        parts.append(f.read_text(encoding="utf-8", errors="replace"))
    else:
        parts.append("(missing)")
SLIDES_SUM.write_text("\n".join(parts), encoding="utf-8")
print("slides summary written")
