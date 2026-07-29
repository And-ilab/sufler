# -*- coding: utf-8 -*-
"""Apply Word comments to operation manual; save result to Desktop only."""
from __future__ import annotations

import copy
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

SRC = Path(
    r"c:\Users\user\Downloads\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina1.docx"
)
OUT = Path(
    r"C:\Users\user\Desktop\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_ispravleno.docx"
)
FIGS = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\manual_figures")
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
BASE_PT = 14

# One representative image per figure (primary for section 1.1.4)
FIGURES_114 = [
    ("Рис. 1. Подключение питания и сжатого воздуха, включение станка", "unpacking_threading_slide01_img1.jpeg"),
    ("Рис. 2. Распаковка станка и снятие опорной рамы", "unpacking_threading_slide02_img1.jpeg"),
    ("Рис. 3. Смещение муфт по осям X и Z (на 2–3 см)", "unpacking_threading_slide03_img1.jpeg"),
    ("Рис. 4. Заправка медной проволоки в акриловые бобины", "unpacking_threading_slide04_img1.jpeg"),
    ("Рис. 5. Проведение провода через ролики", "unpacking_threading_slide05_img1.jpeg"),
    ("Рис. 6. Схема заправки на натяжителе", "unpacking_threading_slide06_img1.jpeg"),
    ("Рис. 7. Провод через натяжитель, антинатяжитель и ролик", "unpacking_threading_slide07_img1.jpeg"),
    ("Рис. 8. Ввод провода в наконечник и зажим ножницами", "unpacking_threading_slide08_img1.jpeg"),
    ("Рис. 9. Режим AUTO и возврат в HOME (нулевое положение)", "unpacking_threading_slide09_img1.jpeg"),
    ("Рис. 10. Ручной режим MAN.: SCISSORS A1/A2 (BACK/OUT)", "unpacking_threading_slide10_img1.jpeg"),
    ("Рис. 11. Автоматический зажим провода кнопками A1 и A2", "unpacking_threading_slide11_img1.jpeg"),
    ("Рис. 12. Фиксация статора (Clamp open / close)", "unpacking_threading_slide12_img1.jpeg"),
    ("Рис. 13. Запуск автоматической намотки (две зелёные кнопки)", "unpacking_threading_slide13_img1.jpeg"),
    ("Рис. 14. Датчики безопасности дверного проёма", "unpacking_threading_slide14_img1.jpeg"),
]

FIGURES_APP_B = [
    ("Рис. 15. Диалоговое окно аварии после включения экрана", "parameter_settings_slide01_img1.jpeg"),
    ("Рис. 16. Переход в меню «Файл»", "parameter_settings_slide02_img1.jpeg"),
    ("Рис. 17. Окно ввода пароля (без ввода — OK)", "parameter_settings_slide03_img1.jpeg"),
    ("Рис. 18. Создание новой программы или вход в режим редактирования", "parameter_settings_slide04_img1.jpeg"),
    ("Рис. 19. Режим редактирования параметров (листание страниц)", "parameter_settings_slide05_img1.jpeg"),
    ("Рис. 20. Вход в «Параметры намотки»", "parameter_settings_slide06_img1.jpeg"),
    ("Рис. 21. Автоматический режим укладки и параметры витков", "parameter_settings_slide07_img1.jpeg"),
    ("Рис. 22. Режим точной укладки по слоям", "parameter_settings_slide08_img1.jpeg"),
    ("Рис. 23. Сохранение и переход в автоматическую намотку", "parameter_settings_slide09_img1.jpeg"),
]


def set_run_font(run, size=BASE_PT, bold=None):
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def force_all_fonts_14(doc: Document):
    """Comment #3: 14 pt everywhere."""
    for p in doc.paragraphs:
        for run in p.runs:
            # keep relative hierarchy: if was clearly a large title (>16), keep 16
            cur = run.font.size.pt if run.font.size else BASE_PT
            if cur >= 18:
                set_run_font(run, size=16, bold=run.bold)
            elif cur >= 15:
                set_run_font(run, size=14, bold=True if run.bold is None else run.bold)
            else:
                set_run_font(run, size=BASE_PT, bold=run.bold)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=BASE_PT, bold=run.bold)


def replace_in_paragraph(p, old: str, new: str) -> bool:
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    # simplest: put whole text into first run
    if not p.runs:
        run = p.add_run(new)
        set_run_font(run)
        return True
    new_full = full.replace(old, new)
    p.runs[0].text = new_full
    set_run_font(p.runs[0], bold=p.runs[0].bold)
    for r in p.runs[1:]:
        r.text = ""
    return True


def replace_everywhere(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, old, new):
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_paragraph(p, old, new):
                        n += 1
    return n


def find_para_index(doc: Document, exact: str | None = None, contains: str | None = None) -> int:
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if exact is not None and t == exact:
            return i
        if contains is not None and contains in t:
            return i
    raise KeyError(f"paragraph not found: exact={exact!r} contains={contains!r}")


def insert_paragraph_after(paragraph, text="", *, size=BASE_PT, bold=False, align="justify", space_after=6, first_line=True):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    # wrap as paragraph
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line and align == "justify" and text:
        pf.first_line_indent = Cm(1.25)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


# Slightly smaller than full text width so figures don't dominate the page
IMG_WIDTH_CM = 10.5


def insert_picture_after(paragraph, image_path: Path, width_cm=IMG_WIDTH_CM):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


def apply_page_setup(doc: Document):
    """A4 margins only — no empty page border / empty stamp on pages without figures."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgBorders"):
            sectPr.remove(child)

    # Keep footer clean (no empty GOST stamp cells)
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    for tbl in list(footer.tables):
        tbl._element.getparent().remove(tbl._element)


def shift_year_down(doc: Document):
    """Comment #0: move '2026' lower on title page."""
    idx = find_para_index(doc, exact="2026")
    p = doc.paragraphs[idx]
    p.paragraph_format.space_before = Pt(180)  # push down
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, size=BASE_PT, bold=False)


def insert_org_block(doc: Document):
    """Comment #2: manufacturer + supplier / RB representative / service center."""
    # After introduction opening paragraphs — insert after first intro para about purpose
    idx = find_para_index(doc, contains="Настоящее руководство по эксплуатации (далее — РЭ) предназначено")
    # find the paragraph about РЭ содержит сведения — insert after the composition block
    # Better: after "Настоящее РЭ распространяется на двухстанционную..."
    try:
        idx = find_para_index(doc, contains="Настоящее РЭ распространяется на двухстанционную намоточную машину")
    except KeyError:
        idx = find_para_index(doc, contains="Виды опасных воздействий")

    anchor = doc.paragraphs[idx]
    blocks = [
        ("Предприятие-изготовитель", True, "center"),
        (
            "GREWIN INDUSTRIAL GROUP CO., LTD. (КНР, г. Тяньцзинь).\n"
            "Адрес: 4th floor, A-A3 zone aviation business area, Dongli Distr, Tianjin, China.\n"
            "Тел./WhatsApp: +8618622096679; e-mail: admin@coilwindingmachinechina.com; "
            "www.coilwindingmachinechina.com.",
            False,
            "justify",
        ),
        (
            "Поставщик, официальный представитель и сервисный центр в Республике Беларусь",
            True,
            "center",
        ),
        (
            "ООО «Финсельват».\n"
            "Адрес: Минская обл., Минский р-н, Новодворский сельсовет, д. Большое Стиклево, "
            "д. 40, к. 2, оф. 52.\n"
            "УНП 692204462; тел.: +375 29 658 53 63; e-mail: finselvat.info@yandex.ru; "
            "www.цифровая.бел.\n"
            "р/с BY21ALFA30122C18740010270000 в ЗАО «Альфа-банк», г. Минск, ул. Сурганова, 43–47, "
            "код ALFABY2X.",
            False,
            "justify",
        ),
    ]
    # insert in reverse so order is preserved
    cursor = anchor
    created = []
    for text, bold, align in blocks:
        # multi-line: split into paragraphs
        parts = text.split("\n")
        for j, part in enumerate(parts):
            p = insert_paragraph_after(
                cursor,
                part,
                size=BASE_PT,
                bold=bold and j == 0,
                align=align if j == 0 and bold else ("justify" if not bold else align),
                space_after=3 if j < len(parts) - 1 else 8,
                first_line=(not bold),
            )
            created.append(p)
            cursor = p
    return created


def insert_figures_after(doc: Document, after_contains: str, figures, heading: str | None):
    idx = find_para_index(doc, contains=after_contains)
    # For 1.1.4: insert after last body paragraph of the section — use the alarm paragraph
    cursor = doc.paragraphs[idx]
    if heading:
        cursor = insert_paragraph_after(
            cursor, heading, size=BASE_PT, bold=True, align="left", space_after=8, first_line=False
        )
    for caption, fname in figures:
        path = FIGS / fname
        if not path.exists():
            raise FileNotFoundError(path)
        cursor = insert_picture_after(cursor, path)
        cursor = insert_paragraph_after(
            cursor, caption, size=BASE_PT, bold=False, align="center", space_after=10, first_line=False
        )


def strip_comments_from_docx(path: Path):
    """Remove comment parts and comment markers from document XML."""
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        namelist = zin.namelist()
        skip = {
            "word/comments.xml",
            "word/commentsExtended.xml",
            "word/commentsExtensible.xml",
            "word/commentsIds.xml",
            "word/people.xml",
        }
        for name in namelist:
            if name in skip:
                continue
            data = zin.read(name)
            if name == "word/document.xml":
                root = ET.fromstring(data)
                # remove commentRangeStart/End and commentReference
                for el in list(root.iter()):
                    for child in list(el):
                        if child.tag in (
                            W + "commentRangeStart",
                            W + "commentRangeEnd",
                            W + "commentReference",
                        ):
                            el.remove(child)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            if name == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                NS = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
                for rel in list(root):
                    target = rel.get("Target", "")
                    if "comment" in target.lower() or "people.xml" in target.lower():
                        root.remove(rel)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            if name == "[Content_Types].xml":
                root = ET.fromstring(data)
                for el in list(root):
                    pn = el.get("PartName", "")
                    if "comment" in pn.lower() or pn.endswith("/people.xml"):
                        root.remove(el)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(name, data)
    tmp.replace(path)



def remove_end_of_document_marker(doc: Document) -> int:
    """Remove trailing 'Конец документа.' paragraph(s) and blank spacers after change sheet."""
    removed = 0
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if t == "Конец документа." or t.startswith("Конец документа"):
            el = p._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                removed += 1

    body = doc.element.body
    children = list(body)
    leaf_idx = None
    for i, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        texts = [n.text or "" for n in child.iter(qn("w:t"))]
        text = "".join(texts).strip()
        if text.startswith("Лист регистрации изменений"):
            leaf_idx = i

    if leaf_idx is not None:
        for child in list(children[leaf_idx + 1 :]):
            if child.tag in (qn("w:tbl"), qn("w:sectPr")):
                continue
            if child.tag == qn("w:p"):
                texts = [n.text or "" for n in child.iter(qn("w:t"))]
                text = "".join(texts).strip()
                if not text:
                    parent = child.getparent()
                    if parent is not None:
                        parent.remove(child)
                        removed += 1
    return removed


def main():
    if not SRC.exists():
        raise SystemExit(f"Source not found: {SRC}")
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    # Drop source trailing "Конец документа." / blank spacers after change sheet
    n_end = remove_end_of_document_marker(doc)
    print(f"Removed end-of-document / trailing blanks: {n_end}")

    # #4 languages: add Russian
    n1 = replace_everywhere(doc, "китайский / английский", "русский / китайский / английский")
    n2 = replace_everywhere(doc, "Китайский / английский", "Русский / китайский / английский")
    print(f"Language replacements: {n1 + n2}")

    # #0 year down
    shift_year_down(doc)
    print("Year shifted down")

    # #2 org block
    insert_org_block(doc)
    print("Org block inserted")

    # #5 figures for 1.1.4 — after the safety-stop paragraph of the section
    insert_figures_after(
        doc,
        after_contains="При срабатывании датчиков безопасности в зоне дверного проёма цикл немедленно останавливается",
        figures=FIGURES_114,
        heading="Иллюстрации к разделу 1.1.4 (устройство, подготовка и работа)",
    )
    print(f"Inserted {len(FIGURES_114)} figures into 1.1.4")

    # Appendix B figures — only in body (before last «Лист регистрации изменений»)
    try:
        paras = [p.text.strip() for p in doc.paragraphs]
        print("Appendix B indices:", [i for i, t in enumerate(paras) if t.startswith("Приложение Б")])
        leaf_idxs = [i for i, t in enumerate(paras) if t.startswith("Лист регистрации")]
        if not leaf_idxs:
            raise RuntimeError("Лист регистрации not found")
        end = leaf_idxs[-1] - 1
        while end > 0 and not doc.paragraphs[end].text.strip():
            end -= 1
        cursor = doc.paragraphs[end]
        cursor = insert_paragraph_after(
            cursor,
            "Иллюстрации к приложению Б (настройка параметров на сенсорном экране)",
            size=BASE_PT,
            bold=True,
            align="left",
            space_after=8,
            first_line=False,
        )
        for caption, fname in FIGURES_APP_B:
            path = FIGS / fname
            if not path.exists():
                raise FileNotFoundError(path)
            cursor = insert_picture_after(cursor, path)
            cursor = insert_paragraph_after(
                cursor, caption, size=BASE_PT, bold=False, align="center", space_after=10, first_line=False
            )
        print(f"Inserted {len(FIGURES_APP_B)} figures into Appendix B")
    except Exception as e:
        print("Appendix B figures error:", e)

    # Page setup without empty borders on pages without figures
    apply_page_setup(doc)
    print("Page setup applied (no empty frame)")

    # #3 font 14 everywhere (after inserts)
    force_all_fonts_14(doc)
    print("Fonts set to 14 pt")

    doc.save(str(OUT))
    strip_comments_from_docx(OUT)
    print(f"Saved: {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
