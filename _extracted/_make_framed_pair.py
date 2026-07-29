# -*- coding: utf-8 -*-
"""From folder 333: keep no-frame file, create twin WITH GOST page frames.

- No empty footer stamp tables
- Page border only (рамка листа)
- Images slightly smaller + spacing so they don't overlap / collide with frame
"""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

FOLDER = Path(r"C:\Users\user\Desktop\333")
SRC_CANDIDATES = [
    FOLDER / "Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_ispravleno.docx",
    FOLDER / "Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_bez_ramok.docx",
]
OUT_BEZ = FOLDER / "Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_bez_ramok.docx"
OUT_S = FOLDER / "Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_s_ramkami.docx"

# Inside frame printable width ~ A4 210 - L25 - R15 = 170 mm → keep images ~95 mm
IMG_WIDTH_CM = 9.5
IMG_SPACE_BEFORE_PT = 10
IMG_SPACE_AFTER_PT = 8
CAPTION_SPACE_AFTER_PT = 14

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def clear_dir_keep_none(folder: Path):
    pass


def find_source() -> Path:
    for p in SRC_CANDIDATES:
        if p.exists():
            return p
    raise FileNotFoundError(f"No source docx in {FOLDER}")


def ensure_bez_ramok(src: Path) -> Path:
    """Keep a clean no-frame copy named *_bez_ramok.docx."""
    if src.resolve() == OUT_BEZ.resolve():
        return OUT_BEZ
    shutil.copy2(src, OUT_BEZ)
    # If source was the old "ispravleno" name, remove it after copying to avoid 3 files
    # (user wants exactly two files: with and without). Only delete if different name.
    if src.name != OUT_BEZ.name and src.name != OUT_S.name:
        try:
            src.unlink()
        except PermissionError:
            print(f"WARNING: could not remove old name {src.name} (locked); leave it")
    return OUT_BEZ


def add_page_frame(doc: Document):
    """ЕСКД-like page frame. No empty stamp footer."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Content inside the frame (border sits near page edge)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgBorders"):
            sectPr.remove(child)

    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    # space = distance from page edge to border line (points)
    # left a bit larger (binding / ЕСКД left field)
    spaces = {"top": "14", "left": "20", "bottom": "14", "right": "14"}
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "24")  # 1.5 pt
        el.set(qn("w:space"), spaces[edge])
        el.set(qn("w:color"), "000000")
        pgBorders.append(el)
    sectPr.append(pgBorders)

    # Clean footer — no empty stamp cells ("пустые рамки")
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    for tbl in list(footer._element):
        if tbl.tag == qn("w:tbl"):
            footer._element.remove(tbl)
    footer.add_paragraph()


def resize_inline_images(doc: Document, width_cm: float = IMG_WIDTH_CM):
    """Force all inline pictures to the same width; keep aspect ratio via cx only
    (Word will keep aspect if we also scale cy proportionally)."""
    target_cx = int(Cm(width_cm))  # EMUs
    body = doc.element.body
    for extent in body.iter(WP + "extent"):
        cx = int(extent.get("cx", "0"))
        cy = int(extent.get("cy", "0"))
        if cx <= 0:
            continue
        # scale height proportionally
        new_cy = int(cy * (target_cx / cx)) if cy > 0 else cy
        extent.set("cx", str(target_cx))
        extent.set("cy", str(new_cy))
        # also update a:ext inside a:xfrm if present
        # parent chain: wp:inline/wp:anchor -> a:graphic -> ...
    # Update drawingml xfrm extents that mirror size
    for ext in body.iter(A + "ext"):
        # only those with both cx and cy that look like image size (large)
        cx = ext.get("cx")
        cy = ext.get("cy")
        if not cx or not cy:
            continue
        try:
            icx, icy = int(cx), int(cy)
        except ValueError:
            continue
        if icx < 100000:  # skip tiny decorative
            continue
        new_cy = int(icy * (target_cx / icx)) if icx else icy
        ext.set("cx", str(target_cx))
        ext.set("cy", str(new_cy))


def space_figure_paragraphs(doc: Document):
    """Add spacing around paragraphs that contain pictures or start with «Рис.»."""
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        has_pic = bool(p._element.findall(".//" + WP + "inline")) or bool(
            p._element.findall(".//" + WP + "anchor")
        )
        is_caption = text.startswith("Рис.")
        if has_pic:
            p.paragraph_format.space_before = Pt(IMG_SPACE_BEFORE_PT)
            p.paragraph_format.space_after = Pt(IMG_SPACE_AFTER_PT)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # ensure inline (not floating overlap): convert anchor-> leave as is if inline
        elif is_caption:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(CAPTION_SPACE_AFTER_PT)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _register_ooxml_prefixes():
    # Keep readable prefixes when ElementTree rewrites document.xml
    ET.register_namespace("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
    ET.register_namespace("r", "http://schemas.openxmlformats.org/officeDocument/2006/relationships")
    ET.register_namespace("wp", "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing")
    ET.register_namespace("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
    ET.register_namespace("pic", "http://schemas.openxmlformats.org/drawingml/2006/picture")
    ET.register_namespace("mc", "http://schemas.openxmlformats.org/markup-compatibility/2006")
    ET.register_namespace("w14", "http://schemas.microsoft.com/office/word/2010/wordml")
    ET.register_namespace("w15", "http://schemas.microsoft.com/office/word/2012/wordml")
    ET.register_namespace("wp14", "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing")


def convert_floating_to_inline(docx_path: Path):
    """If any wp:anchor (floating) images exist, convert to wp:inline to prevent overlap."""
    _register_ooxml_prefixes()
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = ET.fromstring(data)
                # Replace wp:anchor with wp:inline keeping children where possible
                parent_map = {c: p for p in root.iter() for c in p}
                for anchor in list(root.iter(WP + "anchor")):
                    parent = parent_map.get(anchor)
                    if parent is None:
                        continue
                    inline = ET.Element(WP + "inline")
                    # copy dist attributes if any
                    for attr in ("distT", "distB", "distL", "distR"):
                        if attr in anchor.attrib:
                            inline.set(attr, anchor.get(attr))
                    # move children except positioning-only nodes
                    skip = {
                        WP + "simplePos",
                        WP + "positionH",
                        WP + "positionV",
                        WP + "wrapNone",
                        WP + "wrapSquare",
                        WP + "wrapTight",
                        WP + "wrapThrough",
                        WP + "wrapTopAndBottom",
                    }
                    for child in list(anchor):
                        if child.tag in skip:
                            continue
                        inline.append(child)
                    # ensure extent + docPr + graphic exist (already moved)
                    idx = list(parent).index(anchor)
                    parent.remove(anchor)
                    parent.insert(idx, inline)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(name, data)
    tmp.replace(docx_path)


def verify(path: Path) -> dict:
    info = {"path": str(path), "size": path.stat().st_size}
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        info["has_comments"] = "word/comments.xml" in names
        info["media"] = sum(1 for n in names if n.startswith("word/media/"))
        doc = z.read("word/document.xml")
        info["pgBorders"] = b"pgBorders" in doc
        root = ET.fromstring(doc)
        # Count by Clark notation (prefixes may be a:/wp: or nsN: after ET rewrite)
        info["blips"] = sum(1 for _ in root.iter(A + "blip"))
        info["anchors"] = sum(1 for _ in root.iter(WP + "anchor"))
        info["inlines"] = sum(1 for _ in root.iter(WP + "inline"))
        cxs = sorted({e.get("cx") for e in root.iter(WP + "extent") if e.get("cx")})
        info["extent_cx"] = cxs
        # footer stamp check
        footer_inv = False
        for n in names:
            if n.startswith("word/footer") and n.endswith(".xml"):
                ft = z.read(n).decode("utf-8", errors="replace")
                if "Инв" in ft:
                    footer_inv = True
        info["footer_Inv"] = footer_inv
    return info


def main():
    FOLDER.mkdir(parents=True, exist_ok=True)
    src = find_source()
    print("Source:", src)

    bez = ensure_bez_ramok(src)
    print("Bez ramok:", bez, bez.stat().st_size)

    # Build with-frames from bez
    shutil.copy2(bez, OUT_S)
    doc = Document(str(OUT_S))
    add_page_frame(doc)
    resize_inline_images(doc, IMG_WIDTH_CM)
    space_figure_paragraphs(doc)
    doc.save(str(OUT_S))
    convert_floating_to_inline(OUT_S)

    # Re-apply spacing after XML tweak (optional reopen)
    doc2 = Document(str(OUT_S))
    space_figure_paragraphs(doc2)
    resize_inline_images(doc2, IMG_WIDTH_CM)
    # ensure frame still present
    add_page_frame(doc2)
    doc2.save(str(OUT_S))

    # List folder
    print("\nFolder contents:")
    for p in sorted(FOLDER.glob("*")):
        if p.is_file():
            print(f"  {p.name}  ({p.stat().st_size} bytes)")

    print("\nVerify bez:", verify(OUT_BEZ))
    print("Verify s_ramkami:", verify(OUT_S))

    # Exactly two docx preferred — warn if extras
    docxs = list(FOLDER.glob("*.docx"))
    if len(docxs) != 2:
        print(f"WARNING: expected 2 docx, found {len(docxs)}: {[p.name for p in docxs]}")


if __name__ == "__main__":
    main()
