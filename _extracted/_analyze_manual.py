# -*- coding: utf-8 -*-
import zipfile
import re
from pathlib import Path
from xml.etree import ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

DOCX = Path(r"c:\Users\user\Downloads\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina1.docx")
OUT = Path(r"C:\Users\user\Desktop\sufler\sufler\_extracted\manual_structure.txt")

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "v": "urn:schemas-microsoft-com:vml",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

def twips_to_cm(twips):
    if twips is None:
        return None
    try:
        return round(int(twips) / 1440 * 2.54, 3)
    except Exception:
        return twips

def twips_to_pt(twips):
    if twips is None:
        return None
    try:
        return round(int(twips) / 20, 1)
    except Exception:
        return twips

def emu_to_cm(emu):
    try:
        return round(int(emu) / 914400 * 2.54, 3)
    except Exception:
        return None

lines = []
def w(s=""):
    lines.append(s if isinstance(s, str) else str(s))

doc = Document(str(DOCX))

w("=" * 80)
w("MANUAL STRUCTURE ANALYSIS")
w(f"Source: {DOCX}")
w(f"Size bytes: {DOCX.stat().st_size}")
w("=" * 80)

# --- Sections / page setup ---
w("\n## 3. PAGE SETUP / SECTIONS")
w(f"Sections count: {len(doc.sections)}")
for i, sec in enumerate(doc.sections):
    w(f"\n--- Section {i} ---")
    try:
        w(f"  page_width_cm: {round(sec.page_width.cm, 3) if sec.page_width else None}")
        w(f"  page_height_cm: {round(sec.page_height.cm, 3) if sec.page_height else None}")
        w(f"  orientation: {sec.orientation}")
        w(f"  left_margin_cm: {round(sec.left_margin.cm, 3) if sec.left_margin else None}")
        w(f"  right_margin_cm: {round(sec.right_margin.cm, 3) if sec.right_margin else None}")
        w(f"  top_margin_cm: {round(sec.top_margin.cm, 3) if sec.top_margin else None}")
        w(f"  bottom_margin_cm: {round(sec.bottom_margin.cm, 3) if sec.bottom_margin else None}")
        w(f"  header_distance_cm: {round(sec.header_distance.cm, 3) if sec.header_distance else None}")
        w(f"  footer_distance_cm: {round(sec.footer_distance.cm, 3) if sec.footer_distance else None}")
        w(f"  different_first_page_header_footer: {sec.different_first_page_header_footer}")
    except Exception as e:
        w(f"  error: {e}")

# --- Headers / footers ---
w("\n## 4. HEADERS / FOOTERS")
for i, sec in enumerate(doc.sections):
    for kind, part in [
        ("header", sec.header),
        ("footer", sec.footer),
        ("first_page_header", sec.first_page_header),
        ("first_page_footer", sec.first_page_footer),
        ("even_page_header", sec.even_page_header),
        ("even_page_footer", sec.even_page_footer),
    ]:
        try:
            texts = []
            for p in part.paragraphs:
                t = p.text.strip()
                if t:
                    texts.append(t)
            # also tables in header/footer
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        ct = cell.text.strip()
                        if ct:
                            texts.append(f"[table] {ct}")
            linked = getattr(part, "is_linked_to_previous", None)
            w(f"Section {i} {kind}: linked_to_previous={linked}")
            if texts:
                for t in texts:
                    w(f"  | {t}")
            else:
                w("  | (empty)")
        except Exception as e:
            w(f"Section {i} {kind}: error {e}")

# --- Paragraphs with styles / fonts ---
w("\n## 1. FULL DOCUMENT PLAIN TEXT (paragraphs)")
w("(style | approx_font_pt | text)")
w("-" * 80)

def para_font_size_pt(p):
    sizes = []
    for run in p.runs:
        if run.font.size is not None:
            try:
                sizes.append(run.font.size.pt)
            except Exception:
                pass
    # style default
    try:
        if p.style and p.style.font and p.style.font.size:
            sizes.append(p.style.font.size.pt)
    except Exception:
        pass
    if not sizes:
        # try XML
        for rPr in p._element.findall(".//w:rPr/w:sz", NS):
            val = rPr.get(qn("w:val"))
            if val:
                try:
                    sizes.append(int(val) / 2)
                except Exception:
                    pass
        for rPr in p._element.findall(".//w:rPr/w:szCs", NS):
            val = rPr.get(qn("w:val"))
            if val:
                try:
                    sizes.append(int(val) / 2)
                except Exception:
                    pass
    if sizes:
        # most common / first
        return sizes[0]
    return None

# count drawings per paragraph
def count_drawings(p):
    el = p._element
    n = 0
    n += len(el.findall(".//w:drawing", NS))
    n += len(el.findall(".//w:pict", NS))
    n += len(el.findall(".//w:object", NS))
    return n

image_events = []  # (para_index, before_text, after_text, count)
headings = []
all_para_texts = []

for idx, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "(none)"
    fpt = para_font_size_pt(p)
    text = p.text
    # keep all texts; for very long image-like captions don't truncate content away - user asked list ALL
    # but "Truncate image captions carefully" - meaning if caption is huge binary-like, truncate carefully
    display = text
    if len(display) > 2000:
        display = display[:2000] + f"... [truncated, total_len={len(text)}]"
    fstr = f"{fpt:.1f}" if isinstance(fpt, float) else (str(fpt) if fpt else "?")
    w(f"[{idx}] style={style!r} font≈{fstr}pt")
    w(f"    {display}")
    all_para_texts.append((idx, style, text))
    nd = count_drawings(p)
    if nd:
        before = doc.paragraphs[idx - 1].text.strip() if idx > 0 else ""
        after = doc.paragraphs[idx + 1].text.strip() if idx + 1 < len(doc.paragraphs) else ""
        image_events.append((idx, nd, before, after, text.strip()))
    # headings
    if style and (style.startswith("Heading") or style.startswith("Заголовок") or re.match(r"^Title", style) or style in ("Title", "Subtitle", "Название")):
        headings.append((idx, style, text.strip()))
    elif re.match(r"^\d+(\.\d+)*\s+\S", text.strip()):
        # numbered section-like
        if len(text.strip()) < 200:
            headings.append((idx, f"{style} [numbered]", text.strip()))

# tables
w("\n## TABLES")
w(f"Tables count: {len(doc.tables)}")
for ti, table in enumerate(doc.tables):
    w(f"\n--- Table {ti} ({len(table.rows)}x{len(table.columns) if table.rows else 0}) ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " | ").strip() for c in row.cells]
        w(f"  R{ri}: " + " || ".join(cells))

# inline shapes
w("\n## 2. IMAGES / INLINE SHAPES")
inline_shapes = list(doc.inline_shapes)
w(f"python-docx inline_shapes count: {len(inline_shapes)}")
for i, shape in enumerate(inline_shapes):
    try:
        w(f"  inline[{i}]: type={shape.type}, width_cm={round(shape.width.cm,3) if shape.width else None}, height_cm={round(shape.height.cm,3) if shape.height else None}")
    except Exception as e:
        w(f"  inline[{i}]: error {e}")

w(f"\nParagraphs containing drawings/pict: {len(image_events)}")
for idx, nd, before, after, own in image_events:
    w(f"  para[{idx}] drawings={nd}")
    w(f"    own_text: {own[:200]!r}")
    w(f"    before_para: {before[:200]!r}")
    w(f"    after_para: {after[:200]!r}")

# zipfile media count
w("\n### Package media (zipfile)")
with zipfile.ZipFile(DOCX, "r") as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    w(f"word/media files: {len(media)}")
    for n in media:
        info = z.getinfo(n)
        w(f"  {n} ({info.file_size} bytes)")
    # also count drawings in document.xml
    xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    drawings = root.findall(".//w:drawing", NS)
    picts = root.findall(".//w:pict", NS)
    blips = root.findall(".//a:blip", NS)
    w(f"document.xml w:drawing: {len(drawings)}")
    w(f"document.xml w:pict: {len(picts)}")
    w(f"document.xml a:blip: {len(blips)}")

# Title page / 2026
w("\n## 5. TITLE PAGE / '2026' CONTEXT")
for idx, style, text in all_para_texts:
    if "2026" in text or idx < 40:
        if "2026" in text or (idx < 40 and text.strip()):
            marker = " <<<2026>>>" if "2026" in text else ""
            w(f"[{idx}] {style}: {text.strip()[:500]}{marker}")

# Languages
w("\n## 6. LANGUAGE MENTIONS (китайский / английский / русский)")
lang_re = re.compile(r"китайск|английск|русск|Chinese|English|Russian|中文|язык", re.I)
found_lang = False
for idx, style, text in all_para_texts:
    if lang_re.search(text):
        found_lang = True
        w(f"[{idx}] {style}: {text.strip()}")
# also tables
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if lang_re.search(cell.text):
                found_lang = True
                w(f"table[{ti}] R{ri}C{ci}: {cell.text.strip()}")
if not found_lang:
    w("(none found)")

# Section 1.1.4
w("\n## 7. SECTION 1.1.4 AND SURROUNDINGS")
sec_start = None
sec_end = None
for idx, style, text in all_para_texts:
    t = text.strip()
    if re.match(r"^1\.1\.4\b", t) or re.search(r"\b1\.1\.4\b", t):
        if sec_start is None:
            sec_start = idx
    if sec_start is not None and sec_end is None and idx > sec_start:
        # next section at same or higher level like 1.1.5 or 1.2 or 1.2.x or 2.
        if re.match(r"^(1\.1\.[5-9]|1\.[2-9]|[2-9]\.|1\.1\.1[0-9])\b", t):
            sec_end = idx
            break
if sec_start is None:
    # search loosely
    for idx, style, text in all_para_texts:
        if "1.1.4" in text:
            w(f"loose hit [{idx}]: {text.strip()[:300]}")
    w("(section start marker not found by regex; dumping +/- 30 paras around any 1.1.4)")
    for idx, style, text in all_para_texts:
        if "1.1.4" in text:
            lo = max(0, idx - 5)
            hi = min(len(all_para_texts), idx + 40)
            for j in range(lo, hi):
                w(f"[{j}] {all_para_texts[j][1]}: {all_para_texts[j][2].strip()}")
            break
else:
    if sec_end is None:
        sec_end = min(len(all_para_texts), sec_start + 80)
    # also include a bit before
    lo = max(0, sec_start - 3)
    w(f"Section 1.1.4 paras [{sec_start}:{sec_end}] (showing from {lo})")
    for j in range(lo, sec_end):
        w(f"[{j}] {all_para_texts[j][1]}: {all_para_texts[j][2].strip()}")

# Manufacturer / ООО
w("\n## 8. MANUFACTURER / SUPPLIER / ООО MENTIONS")
mfg_re = re.compile(r"ООО|изготовител|производител|поставщик|manufacturer|supplier|Finselvat|Финсельват|компани|завод", re.I)
found_m = False
for idx, style, text in all_para_texts:
    if mfg_re.search(text):
        found_m = True
        w(f"[{idx}] {style}: {text.strip()}")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        celltext = " | ".join(c.text.strip() for c in row.cells)
        if mfg_re.search(celltext):
            found_m = True
            w(f"table[{ti}] R{ri}: {celltext}")
# headers/footers already listed; scan again
for i, sec in enumerate(doc.sections):
    for kind, part in [("header", sec.header), ("footer", sec.footer)]:
        for p in part.paragraphs:
            if mfg_re.search(p.text):
                found_m = True
                w(f"section{i}.{kind}: {p.text.strip()}")
if not found_m:
    w("(none found)")

# Headings list
w("\n## 9. LIST OF HEADINGS")
# dedupe while preserving order
seen = set()
for item in headings:
    key = (item[0], item[2])
    if key in seen:
        continue
    seen.add(key)
    w(f"[{item[0]}] {item[1]}: {item[2]}")

# Also scan for Heading styles more carefully + outline
w("\n### Additional: styles used")
style_counts = {}
for idx, style, text in all_para_texts:
    style_counts[style] = style_counts.get(style, 0) + 1
for s, c in sorted(style_counts.items(), key=lambda x: (-x[1], x[0])):
    w(f"  {s}: {c}")

OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")
print(f"WROTE {OUT} ({OUT.stat().st_size} bytes, {len(lines)} lines)")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(inline_shapes)} image_paras={len(image_events)}")
print(f"headings_listed={len(seen)}")
