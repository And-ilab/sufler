import shutil
from pathlib import Path
from docx import Document

src = Path(r"C:\Users\user\Desktop\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_ispravleno.docx")
dst = Path(r"C:\Users\user\Desktop\Rukovodstvo_po_ekspluatacii_dvuhstancionnaya_namotochnaya_mashina_ispravleno2.docx")

shutil.copy2(src, dst)
print(f"COPIED: {dst}")

def clean_document(path: Path) -> None:
    doc = Document(str(path))

    to_remove = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t == "Конец документа." or t.startswith("Конец документа"):
            to_remove.append(p)

    for p in to_remove:
        text_preview = (p.text or "")[:80]
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
        print(f"REMOVED paragraph: {text_preview!r}")

    while True:
        paras = list(doc.paragraphs)
        last_reg_idx = None
        for i, p in enumerate(paras):
            if "Лист регистрации изменений" in (p.text or ""):
                last_reg_idx = i
        if last_reg_idx is None:
            break
        if len(paras) <= last_reg_idx + 1:
            break
        last = paras[-1]
        if (last.text or "").strip():
            break
        if paras.index(last) > last_reg_idx:
            el = last._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                print("REMOVED trailing empty paragraph")
        else:
            break

    doc.save(str(path))
    print(f"SAVED: {path}")

clean_document(dst)

original_ok = False
try:
    clean_document(src)
    original_ok = True
except PermissionError as e:
    print(f"ORIGINAL Permission denied: {e}")
    print("USER should close Word and use ispravleno2.docx")
except OSError as e:
    print(f"ORIGINAL OSError: {e}")
    print("USER should close Word and use ispravleno2.docx")
except Exception as e:
    print(f"ORIGINAL failed: {type(e).__name__}: {e}")
    print("USER should close Word and use ispravleno2.docx")

saved = src if original_ok else dst
print(f"\n=== Last 5 non-empty paragraphs of: {saved} ===")
doc = Document(str(saved))
nonempty = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
for t in nonempty[-5:]:
    print(repr(t[:200] if len(t) > 200 else t))

print("\n=== SUCCEEDED PATHS ===")
print(f"copy cleaned: {dst}")
if original_ok:
    print(f"original cleaned: {src}")
else:
    print("original: FAILED (leave as-is)")