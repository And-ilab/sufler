# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).with_name("Otvety-New-IT-Outsource-SUZ-RAG.docx")

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def set_run_font(run, bold=False, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading_ru(text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        size = 14
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        size = 12
    set_run_font(p.add_run(text), bold=True, size=size)
    return p


def add_p(text, bold=False):
    p = doc.add_paragraph()
    set_run_font(p.add_run(text), bold=bold, size=12)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    set_run_font(p.add_run(text), size=12)
    return p


def add_num(text):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    set_run_font(p.add_run(text), size=12)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_run_font(p.add_run(text), bold=bold, size=10)


def shade_header_cell(cell, color="D9E2F3"):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tc.append(shd)


# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    t.add_run(
        "Ответы на уточняющие вопросы по интеграции СУЗ (1С-Битрикс) и RAG"
    ),
    bold=True,
    size=16,
)

for m in [
    "Адресат: Нью Айти Аутсорс",
    "Тема: уточнение информации для оценки работ по интеграции СУЗ и RAG",
    "Дата: 23.07.2026",
    "Основание: согласованное ТЗ интеграции (модель B — исходящий webhook с полным текстом статьи)",
]:
    p = doc.add_paragraph()
    set_run_font(p.add_run(m), size=11)

doc.add_paragraph()
add_p("Добрый день.")
add_p(
    "Ниже ответы на вопросы по интеграции СУЗ (1С-Битрикс) и RAG. "
    "Опираемся на согласованное ТЗ интеграции (модель B: исходящий webhook с полным текстом статьи)."
)
add_p(
    "Срок. Ответы направляем сейчас, 23.07.2026. Отдельные фактические параметры стенда СУЗ "
    "(точный ID инфоблока, полный список кодов свойств, точная схема workflow) в ТЗ как готовые "
    "значения банка не зафиксированы — их нужно подтвердить владельцем СУЗ / ДИТ. Как только эти "
    "данные будут предоставлены, маппинг можно закрыть окончательно. Ориентир по сроку уточнения "
    "со стороны банка: до конца рабочего дня 24.07.2026 либо в срок, который укажет владелец СУЗ."
)

# 1
add_heading_ru("1. Текущая реализация Базы знаний")
add_p(
    "База знаний СУЗ ведётся в 1С-Битрикс (Система управления сайтом) как инфоблок "
    "со статьями (элементы инфоблока)."
)
add_p("По As-Is из ТЗ:")
for b in [
    "около 1732 статей;",
    "правки идут регулярно, ориентировочно 10–30 изменений в день;",
    "структура — разделы и подразделы, вложенность до 5 уровней;",
    "версионность есть (история изменений);",
    "исходящего webhook в суфлёр сейчас нет — это доработка.",
]:
    add_bullet(b)
add_p(
    "Конкретный ID инфоблока в ТЗ не задан (в примерах условный). Его должен указать владелец СУЗ."
)
add_p(
    "Отдельного свойства «доступно операторам КЦ» нет. Доступность определяется разделами и правами групп. "
    "Черновик, по уточнению со стороны СУЗ, скорее всего связан с активностью статьи (галка активности), "
    "а не с отдельным статусным полем."
)
add_p(
    "Workflow / бизнес-процесс публикации как отдельный банковский регламент в ТЗ детально не описан. "
    "Для интеграции важно не UI-кнопка, а итоговое состояние статьи после сохранения и правила, "
    "по которым модуль sync определяет: черновик, опубликовано или снято с публикации."
)
add_p(
    "Полный перечень кодов свойств инфоблока и финальная схема разделов КЦ фиксируются в регламенте "
    "маппинга (BTX-5) после подтверждения владельцем СУЗ."
)

# 2
add_heading_ru("2. Маппинг полей payload")
add_p(
    "Webhook уходит методом POST на внутренний endpoint суфлёра. "
    "В теле передаются метаданные и полный текст статьи."
)
add_p("Прямое соответствие:", bold=True)
for b in [
    "ID элемента → article_id",
    "IBLOCK_ID → iblock_id",
    "раздел (IBLOCK_SECTION_ID) → section_id",
    "NAME → title",
    "PREVIEW_TEXT → preview",
    "DETAIL_TEXT → body_html и body_plain (plain — текст без HTML)",
]:
    add_bullet(b)
add_p("Вычисляемые поля (формирует модуль интеграции на стороне CMS):", bold=True)
for b in [
    "event_id — новый UUID на каждое событие;",
    "event_type — тип события (публикация, черновик, снятие, удаление и т.д.);",
    "occurred_at — время события;",
    "status — draft / published / archived (не сырое поле Bitrix, а нормализованный статус);",
    'visibility_scope — область видимости для индекса КЦ, например ["kc_operator"];',
    "version_id, version_number, is_current — версия статьи;",
    "checksum — sha256 от нормализованного body_plain;",
    "changed_fields — список изменившихся полей при редактировании;",
    "permalink — стабильная ссылка на статью для оператора;",
    "locale — язык, ru или en.",
]:
    add_bullet(b)

# 3
add_heading_ru("3. Когда должна уходить отправка webhook")
add_p("Обработчики событий Bitrix: создание, изменение и удаление элемента инфоблока.")
add_p("Сценарии:", bold=True)
for b in [
    "создание / первая публикация статьи для КЦ → article.version_published;",
    "сохранение черновика без публикации → article.updated, status=draft;",
    "публикация новой версии → article.version_published (новый version_id);",
    "снятие с публикации → article.unpublished;",
    "удаление статьи → article.deleted;",
    "статья вышла из области КЦ (смена раздела/прав) → publish не отправляем либо отправляем unpublished;",
    "при сбое доставки — повтор из очереди с тем же event_id;",
    "при необходимости — ручной replay из журнала модуля;",
    "первичная загрузка при внедрении — через REST (не webhook);",
    "догрузка пропусков — через API /changes.",
]:
    add_bullet(b)

# 4
add_heading_ru("4. Правила статусов draft / published / archived")
add_p("Отдельного поля status в СУЗ нет. Его вычисляет модуль интеграции.")
for b in [
    "draft — черновик; событие может уйти для аудита, но в рабочий индекс RAG статья не попадает;",
    "published — статья активна и доступна для области КЦ; индексируется;",
    "archived — снята с публикации / архив; из поиска суфлёра убирается (soft delete).",
]:
    add_bullet(b)
add_p(
    "Сырой признак ACTIVE в Bitrix сам по себе не равен published. "
    "Опубликованной для RAG считается только статья, которая проходит правила банка: "
    "активность + нужный раздел/scope КЦ."
)
add_p("Удаление — отдельное событие article.deleted (полное удаление из индекса).")

# 5
add_heading_ru("5. Правила формирования отдельных полей")
for b in [
    'visibility_scope — считается модулем по разделу статьи, whitelist разделов КЦ и правам. '
    'Отдельной галки «для операторов» нет. Для КЦ ожидаемо значение вроде ["kc_operator"]. '
    "Если статья вне scope КЦ — publish не отправляем.",
    "version_id — идентификатор редакции (из истории / логики модуля). При новой публикации — новый id.",
    "version_number — порядковый номер версии (1, 2, 3…).",
    "checksum — sha256 от нормализованного текста body_plain.",
    "changed_fields — при редактировании список изменённых полей (NAME, DETAIL_TEXT и т.п.).",
    "permalink — постоянная ссылка на статью для оператора. При правке текста обычно не меняется.",
    "locale — язык статьи (ru/en). Конкретный источник в CMS уточняется в регламенте маппинга.",
]:
    add_bullet(b)

# 6
add_heading_ru("6. Таблица соответствия полей")
headers = ["Поле payload", "Откуда в Bitrix", "Обязательно", "Формат / правило"]
rows = [
    ["event_id", "модуль sync", "да", "UUID"],
    ["event_type", "логика события", "да", "строка события"],
    ["occurred_at", "время события", "да", "ISO-8601"],
    ["article_id", "ID", "да", "число"],
    ["iblock_id", "IBLOCK_ID", "да", "число"],
    ["section_id", "раздел", "нет", "число"],
    [
        "version_id / version_number / is_current",
        "история + модуль",
        "да для версионных событий",
        "id / номер / признак текущей",
    ],
    ["status", "ACTIVE + правила банка", "да", "draft / published / archived"],
    ["title", "NAME", "да", "строка"],
    ["preview", "PREVIEW_TEXT", "да", "строка"],
    ["body_html / body_plain", "DETAIL_TEXT", "да при публикации", "HTML и текст"],
    ["permalink", "URL статьи", "да", "стабильный URL"],
    ["locale", "язык/настройка CMS", "да", "ru / en"],
    ["visibility_scope", "раздел + права + whitelist", "да", "массив строк"],
    ["checksum", "из body_plain", "да", "sha256"],
    ["changed_fields", "diff при Update", "нет", "массив кодов полей"],
]
table = doc.add_table(rows=1 + len(rows), cols=4)
table.style = "Table Grid"
for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True)
    shade_header_cell(table.rows[0].cells[i])
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        set_cell_text(table.rows[ri + 1].cells[ci], val)
add_p(
    "Для удаления и частичного снятия с публикации полный текст статьи не обязателен."
)

# 7
add_heading_ru("7. Что уже есть и что нужно реализовать на стороне CMS")
add_p("Уже есть в платформе Bitrix:", bold=True)
for b in [
    "инфоблоки, разделы, свойства, активность;",
    "события создания/изменения/удаления;",
    "REST для чтения статей;",
    "агенты по расписанию.",
]:
    add_bullet(b)
add_p("Нужно реализовать модулем интеграции (bank.sufler.sync):", bold=True)
for b in [
    "исходящий webhook;",
    "очередь Outbox;",
    "повторную отправку (Retry);",
    "журнал отправок;",
    "ручной Replay;",
    "API /changes;",
    "расчёт status, visibility_scope, version_*, checksum, changed_fields;",
    "документацию маппинга свойств.",
]:
    add_bullet(b)
add_p(
    "RAG, эмбеддинги, LLM и интерфейс оператора — не зона работ Bitrix, это сторона суфлёра."
)
add_p(
    "Ориентир по объёму доработки CMS в ТЗ: примерно 2–4 недели одного разработчика Bitrix "
    "при готовых доступах и тестовом инфоблоке."
)

add_heading_ru("Как в итоге работает логика")
add_p(
    "Редактор меняет статью в СУЗ → срабатывает событие Bitrix → модуль собирает JSON → "
    "сначала пишет в очередь Outbox → затем отправляет webhook в суфлёр → "
    "суфлёр индексирует статью из тела запроса."
)
add_p(
    "Черновики в рабочий индекс не попадают. Публикация и новая версия — индексируются. "
    "Снятие с публикации — скрывается. Удаление — полностью убирается из индекса."
)

add_heading_ru("Что нужно от банка для финальной оценки «в ноль»")
for b in [
    "Точный IBLOCK_ID (тест и прод).",
    "Список разделов КЦ, которые входят в индекс.",
    "Подтверждение правила: как отличить черновик / опубликовано / архив.",
    "Есть ли workflow/БП публикации сверх истории элемента.",
    "Как определяется язык статьи (locale).",
    "Как строится постоянная ссылка (permalink).",
    "URL webhook для теста/прода и требования ИБ к подписи канала.",
]:
    add_num(b)
add_p(
    "Если нужно, готовы отдельно выслать краткую таблицу только по обязательным полям payload "
    "или созвониться с владельцем СУЗ и закрыть открытые пункты."
)
add_p("С уважением.")

doc.save(OUT)
print(f"OK: {OUT}")
